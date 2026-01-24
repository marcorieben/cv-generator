"""
Module description

Purpose: analyzed as source_code
Expected Lifetime: permanent
Category: SOURCE_CODE
Created: 2025-12-11
Last Updated: 2026-01-24
"""
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Handle imports for both direct execution and module import
try:
    from dialogs import show_warning, select_json_file
except ImportError:
    from scripts.dialogs import show_warning, select_json_file

try:
    from scripts.utils.translations import load_translations, get_text
except ImportError:
    from utils.translations import load_translations, get_text

# Globale Konstante für fehlende Daten
MISSING_DATA_MARKER = "! bitte prüfen !"


def validate_json_structure(data, language="de"):
    """Validiert die Struktur der JSON-Daten und gibt kritische Fehler und Info zurück"""
    translations = load_translations()
    critical = []
    info = []
    
    # Erforderliche Top-Level-Felder
    required_fields = [
        "Vorname", "Nachname", "Hauptrolle", "Nationalität", "Ausbildung",
        "Kurzprofil", "Fachwissen_und_Schwerpunkte", "Aus_und_Weiterbildung",
        "Trainings_und_Zertifizierungen", "Sprachen", "Ausgewählte_Referenzprojekte"
    ]
    
    # Prüfe Arrays
    array_fields = [
        "Fachwissen_und_Schwerpunkte", "Aus_und_Weiterbildung",
        "Trainings_und_Zertifizierungen", "Sprachen", "Ausgewählte_Referenzprojekte"
    ]
    
    for field in required_fields:
        if field not in data:
            critical.append(f"{get_text(translations, 'validation', 'missing_field', language)}: {field}")
        elif data[field] is None:
            critical.append(f"{get_text(translations, 'validation', 'field_is_none', language)}: {field}")
        else:
            # Typ-Prüfungen
            if field == "Hauptrolle":
                if not isinstance(data[field], dict) or "Titel" not in data[field] or "Beschreibung" not in data[field]:
                    critical.append(get_text(translations, 'validation', 'hauptrolle_invalid', language))
                else:
                    beschreibung = data[field]["Beschreibung"]
                    word_count = len(beschreibung.split())
                    if word_count < 5 or word_count > 10:
                        info.append(f"{get_text(translations, 'validation', 'hauptrolle_desc_length', language)} (aktuell {word_count})")
            elif field in ["Vorname", "Nachname", "Nationalität", "Ausbildung", "Kurzprofil"]:
                if not isinstance(data[field], str):
                    info.append(f"Feld '{field}' {get_text(translations, 'validation', 'should_be_string', language)} (ist {type(data[field]).__name__})")
            elif field in array_fields:
                if not isinstance(data[field], list):
                    critical.append(f"Feld '{field}' {get_text(translations, 'validation', 'must_be_array', language)}")
    
    # Entferne die doppelte Definition
    # array_fields = [...
    
    # Prüfe Fachwissen_und_Schwerpunkte Struktur
    if "Fachwissen_und_Schwerpunkte" in data and isinstance(data["Fachwissen_und_Schwerpunkte"], list):
        for i, item in enumerate(data["Fachwissen_und_Schwerpunkte"]):
            if not isinstance(item, dict):
                critical.append(f"Fachwissen_und_Schwerpunkte[{i}]: {get_text(translations, 'validation', 'must_be_object', language)}")
                continue
            if "Kategorie" not in item:
                critical.append(f"Fachwissen_und_Schwerpunkte[{i}]: {get_text(translations, 'validation', 'missing_field', language)} 'Kategorie'")
            if "Inhalt" not in item or not isinstance(item["Inhalt"], list):
                critical.append(f"Fachwissen_und_Schwerpunkte[{i}]: 'Inhalt' {get_text(translations, 'validation', 'must_be_array', language)}")
    
    # Prüfe Aus_und_Weiterbildung Struktur
    if "Aus_und_Weiterbildung" in data and isinstance(data["Aus_und_Weiterbildung"], list):
        for i, item in enumerate(data["Aus_und_Weiterbildung"]):
            if not isinstance(item, dict):
                critical.append(f"Aus_und_Weiterbildung[{i}]: {get_text(translations, 'validation', 'must_be_object', language)}")
                continue
            required = ["Zeitraum", "Institution", "Abschluss"]
            for req in required:
                if req not in item:
                    critical.append(f"Aus_und_Weiterbildung[{i}]: {get_text(translations, 'validation', 'missing_field', language)} '{req}'")
    
    # Prüfe Trainings_und_Zertifizierungen Struktur
    if "Trainings_und_Zertifizierungen" in data and isinstance(data["Trainings_und_Zertifizierungen"], list):
        for i, item in enumerate(data["Trainings_und_Zertifizierungen"]):
            if not isinstance(item, dict):
                critical.append(f"Trainings_und_Zertifizierungen[{i}]: {get_text(translations, 'validation', 'must_be_object', language)}")
                continue
            required = ["Zeitraum", "Institution", "Titel"]
            for req in required:
                if req not in item:
                    critical.append(f"Trainings_und_Zertifizierungen[{i}]: {get_text(translations, 'validation', 'missing_field', language)} '{req}'")
    
    # Prüfe Sprachen Struktur
    if "Sprachen" in data and isinstance(data["Sprachen"], list):
        for i, item in enumerate(data["Sprachen"]):
            if not isinstance(item, dict):
                critical.append(f"Sprachen[{i}]: {get_text(translations, 'validation', 'must_be_object', language)}")
                continue
            if "Sprache" not in item:
                critical.append(f"Sprachen[{i}]: {get_text(translations, 'validation', 'missing_field', language)} 'Sprache'")
            if "Level" not in item or not is_valid_level(item["Level"]):
                info.append(f"Sprachen[{i}]: {get_text(translations, 'validation', 'invalid_level', language)}")
    
    # Prüfe Referenzprojekte Struktur
    if "Ausgewählte_Referenzprojekte" in data and isinstance(data["Ausgewählte_Referenzprojekte"], list):
        for i, item in enumerate(data["Ausgewählte_Referenzprojekte"]):
            if not isinstance(item, dict):
                critical.append(f"Ausgewählte_Referenzprojekte[{i}]: {get_text(translations, 'validation', 'must_be_object', language)}")
                continue
            required = ["Zeitraum", "Rolle", "Kunde", "Tätigkeiten"]
            for req in required:
                if req not in item:
                    critical.append(f"Ausgewählte_Referenzprojekte[{i}]: {get_text(translations, 'validation', 'missing_field', language)} '{req}'")
            if "Tätigkeiten" in item and not isinstance(item["Tätigkeiten"], list):
                critical.append(f"Ausgewählte_Referenzprojekte[{i}]: 'Tätigkeiten' {get_text(translations, 'validation', 'must_be_array', language)}")
    
    return critical, info


def parse_level(level):
    """Parst das Level zu einer Zahl 1-5"""
    if isinstance(level, int):
        return level
    if isinstance(level, str):
        import re
        level = level.strip()
        # Versuche Zahl am Anfang
        match = re.match(r'(\d+)', level)
        if match:
            return int(match.group(1))
            
        # Mappe Text (suche in allen Sprachen)
        translations = load_translations()
        text_to_num = {}
        for i in range(1, 6):
            key = f"level_{i}"
            lvl_trans = translations.get("levels", {}).get(key, {})
            for lang_code in lvl_trans:
                text_to_num[lvl_trans[lang_code].lower()] = i
        
        # Fallback Hardcoded für Sicherheit
        text_to_num.update({
            "muttersprache": 5, "native": 5, "maternelle": 5,
            "verhandlungssicher": 4, "proficient": 4, "courant": 4,
            "sehr gute kenntnisse": 3, "advanced": 3, "avancé": 3,
            "gute kenntnisse": 2, "intermediate": 2, "intermédiaire": 2,
            "grundkenntnisse": 1, "basic": 1, "débutant": 1
        })
        
        return text_to_num.get(level.lower(), 0)
    return 0


def is_valid_level(level):
    """Prüft, ob das Level gültig ist"""
    if isinstance(level, int) and 1 <= level <= 5:
        return True
    if isinstance(level, str):
        import re
        level = level.strip()
        # Zahl am Anfang ist okay
        if re.match(r'^\d+', level):
            val = parse_level(level)
            return 1 <= val <= 5
            
        # Prüfe gegen alle Übersetzungen
        translations = load_translations()
        valid_texts = []
        for i in range(1, 6):
            key = f"level_{i}"
            lvl_trans = translations.get("levels", {}).get(key, {})
            for lang_code in lvl_trans:
                valid_texts.append(lvl_trans[lang_code].lower())
        
        if level.lower() in valid_texts:
            return True
            
        # Fallback Hardcoded
        if level.lower() in ["muttersprache", "verhandlungssicher", "sehr gute kenntnisse", "gute kenntnisse", "grundkenntnisse",
                            "native speaker", "full professional proficiency", "very good knowledge", "good knowledge", "basic knowledge",
                            "langue maternelle", "capacité professionnelle complète", "très bonnes connaissances", "bonnes connaissances", "connaissances de base"]:
            return True
            
    return False


def abs_path(relative_path):
    """Gibt den absoluten Pfad relativ zum Skript-Verzeichnis zurück"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(script_dir, relative_path)


def remove_cell_borders(cell):
    """Entfernt alle Rahmen einer Tabellenzelle mittels XML-Manipulation und setzt Padding auf 0"""
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
    tcPr = cell._element.get_or_add_tcPr()
    
    # Rahmen entfernen
    tcBorders = parse_xml(r'<w:tcBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'</w:tcBorders>' % nsdecls('w'))
    # Entferne existierende Rahmen
    existing = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

    # Padding auf 0 setzen
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_padding_zero(cell):
    """Setzt alle internen Zellenabstände (Padding) auf 0"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


# -------------------------------------
# Stil-Funktionen für Formatierung
# -------------------------------------


# Styles aus JSON laden
def load_styles(filename="styles.json"):
    path = abs_path(filename)
    with open(path, "r", encoding="utf-8") as sf:
        return json.load(sf)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    s = styles["heading1"]
    p.paragraph_format.space_before = Pt(s["space_before"])
    p.paragraph_format.space_after = Pt(s["space_after"])
    
    run = p.add_run(text)
    run.font.name = s["font"]
    run.font.size = Pt(s["size"])
    run.font.bold = s["bold"]
    c = s["color"]
    run.font.color.rgb = RGBColor(c[0], c[1], c[2])
    return p


def add_heading_2(doc, text, bold=None):
    p = doc.add_paragraph()
    s = styles["heading2"]
    p.paragraph_format.space_before = Pt(s["space_before"])
    p.paragraph_format.space_after = Pt(s["space_after"])

    run = p.add_run(text)
    run.font.name = s["font"]
    run.font.size = Pt(s["size"])
    run.font.bold = bold if bold is not None else s["bold"]
    c = s["color"]
    run.font.color.rgb = RGBColor(c[0], c[1], c[2])
    return p


def add_normal_text(doc, text):
    p = doc.add_paragraph()
    s = styles["text"]
    p.paragraph_format.space_before = Pt(s.get("space_before", 0))
    p.paragraph_format.space_after = Pt(s.get("space_after", 6))
    add_text_with_highlight(p, text, s["font"], s["size"], s["color"])
    return p


def add_text_with_highlight(paragraph, text, font_name, font_size, font_color, bold=False, italic=False):
    """Helper function to add text (highlighting wird am Ende global durchgeführt)"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*font_color) if isinstance(font_color, (list, tuple)) else font_color
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True


def highlight_missing_data_in_document(doc, translations=None, language="de"):
    """
    Durchsucht das gesamte Dokument nach verschiedenen Varianten des Fehlenden-Daten-Markers,
    normalisiert sie auf die einheitliche Version und hebt alle Vorkommen gelb hervor.
    """
    from docx.enum.text import WD_COLOR_INDEX
    
    if translations is None:
        translations = load_translations()
    
    marker_base = get_text(translations, "system", "missing_data_marker", language)
    
    # Alle möglichen Varianten des Markers (OpenAI verwendet oft andere Bindestriche/Leerzeichen)
    marker_variants = [
        marker_base,
        marker_base.replace(" !", "!"), # Ohne Leerzeichen vor !
        "! bitte prüfen!",           # Fallback DE
        "! bitte prüfen !",          # Fallback DE
        "! please check !",          # Fallback EN
        "! please check!",           # Fallback EN
        "! à vérifier !",            # Fallback FR
        "! à vérifier!",             # Fallback FR
    ]
    
    # Durchsuche alle Paragraphen im Hauptdokument
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Prüfe alle Varianten
            found_variant = None
            for variant in marker_variants:
                if variant in run.text:
                    found_variant = variant
                    break
            
            if found_variant:
                # Split den Run-Text am gefundenen Marker
                parts = run.text.split(found_variant)
                
                # Lösche den aktuellen Run-Text
                run.text = ""
                
                # Füge Teile und normalisierte Marker hinzu
                for i, part in enumerate(parts):
                    if i > 0:
                        # Füge neuen Run mit normalisiertem Marker und Highlighting hinzu
                        new_run = paragraph.add_run(MISSING_DATA_MARKER)
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb
                        new_run.font.bold = run.font.bold
                        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                    if part:
                        # Füge Text-Teil hinzu
                        text_run = paragraph.add_run(part)
                        text_run.font.name = run.font.name
                        text_run.font.size = run.font.size
                        text_run.font.color.rgb = run.font.color.rgb
                        text_run.font.bold = run.font.bold
    
    # Durchsuche auch Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Prüfe alle Varianten
                        found_variant = None
                        for variant in marker_variants:
                            if variant in run.text:
                                found_variant = variant
                                break
                        
                        if found_variant:
                            parts = run.text.split(found_variant)
                            run.text = ""
                            
                            for i, part in enumerate(parts):
                                if i > 0:
                                    new_run = paragraph.add_run(MISSING_DATA_MARKER)
                                    new_run.font.name = run.font.name
                                    new_run.font.size = run.font.size
                                    new_run.font.color.rgb = run.font.color.rgb
                                    new_run.font.bold = run.font.bold
                                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                
                                if part:
                                    text_run = paragraph.add_run(part)
                                    text_run.font.name = run.font.name
                                    text_run.font.size = run.font.size
                                    text_run.font.color.rgb = run.font.color.rgb
                                    text_run.font.bold = run.font.bold

                                    text_run.font.bold = run.font.bold


def get_available_width(doc):
    """Calculate available page width in inches based on margins"""
    from docx.shared import Inches, Cm
    section = doc.sections[0]
    page_width = Inches(8.5)  # Standard US Letter width
    left_margin = section.left_margin
    right_margin = section.right_margin
    return page_width - left_margin - right_margin


def add_basic_info_table(doc, hauptrolle_desc, nationalität, ausbildung, language="de", translations=None):
    """
    Render basic info as a 3-column borderless table with white background.
    Column 1 (20%): Labels (bold)
    Column 2 (60%): Values
    Column 3 (20%): Image placeholder (merged across all 3 rows)
    """
    from docx.shared import Inches
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls
    
    table = doc.add_table(rows=3, cols=3)
    # Don't set a style that may not exist; instead, manually remove borders below

    # Remove all table borders completely
    tbl = table._element
    tblPr = tbl.tblPr
    
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    
    # Create and add borders element with no borders
    tblBorders = parse_xml(r'<w:tBorders %s><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tBorders>' % nsdecls('w'))
    
    # Remove existing borders if any
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)

    # Set column widths: 20%, 60%, 20%
    available_width = get_available_width(doc)
    widths = [available_width * 0.20, available_width * 0.50, available_width * 0.30]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # Helper function to set cell background to white and remove individual cell borders
    def set_cell_background_and_borders(cell, fill_color='FFFFFF'):
        """Set cell background color, remove borders and set padding to 0"""
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', fill_color)
        tcPr.append(shd)
        
        # Remove cell borders
        tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tcBorders>' % nsdecls('w'))
        tcPr.append(tcBorders)

        # Set padding to 0
        from docx.oxml.ns import qn
        tcMar = OxmlElement('w:tcMar')
        for margin in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), '0')
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    s = styles["text"]
    
    # Row 1: Hauptrolle
    cell_label = table.rows[0].cells[0]
    cell_value = table.rows[0].cells[1]
    cell_image = table.rows[0].cells[2]
    
    set_cell_background_and_borders(cell_label, 'FFFFFF')
    set_cell_background_and_borders(cell_value, 'FFFFFF')
    set_cell_background_and_borders(cell_image, 'FFFFFF')
    
    cell_label.text = ""
    cell_value.text = ""
    cell_image.text = ""
    
    p_label = cell_label.paragraphs[0]
    p_label.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    label_text = get_text(translations, 'cv', 'role_label', language)
    run_label = p_label.add_run(label_text)
    run_label.font.name = s["font"]
    run_label.font.size = Pt(s["size"])
    run_label.font.bold = True
    run_label.font.color.rgb = RGBColor(*s["color"])
    
    p_value = cell_value.paragraphs[0]
    p_value.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_text_with_highlight(p_value, hauptrolle_desc, s["font"], s["size"], s["color"])
    
    # Row 2: Nationalität
    cell_label2 = table.rows[1].cells[0]
    cell_value2 = table.rows[1].cells[1]
    cell_image2 = table.rows[1].cells[2]
    
    set_cell_background_and_borders(cell_label2, 'FFFFFF')
    set_cell_background_and_borders(cell_value2, 'FFFFFF')
    set_cell_background_and_borders(cell_image2, 'FFFFFF')
    
    cell_label2.text = ""
    cell_value2.text = ""
    cell_image2.text = ""
    
    p_label2 = cell_label2.paragraphs[0]
    p_label2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    label_text2 = get_text(translations, 'cv', 'nationality_label', language)
    run_label2 = p_label2.add_run(label_text2)
    run_label2.font.name = s["font"]
    run_label2.font.size = Pt(s["size"])
    run_label2.font.bold = True
    run_label2.font.color.rgb = RGBColor(*s["color"])
    
    p_value2 = cell_value2.paragraphs[0]
    p_value2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_text_with_highlight(p_value2, nationalität, s["font"], s["size"], s["color"])
    
    # Row 3: Ausbildung
    cell_label3 = table.rows[2].cells[0]
    cell_value3 = table.rows[2].cells[1]
    cell_image3 = table.rows[2].cells[2]
    
    set_cell_background_and_borders(cell_label3, 'FFFFFF')
    set_cell_background_and_borders(cell_value3, 'FFFFFF')
    set_cell_background_and_borders(cell_image3, 'FFFFFF')
    
    cell_label3.text = ""
    cell_value3.text = ""
    cell_image3.text = ""
    
    p_label3 = cell_label3.paragraphs[0]
    p_label3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    label_text3 = get_text(translations, 'cv', 'education_label', language)
    run_label3 = p_label3.add_run(label_text3)
    run_label3.font.name = s["font"]
    run_label3.font.size = Pt(s["size"])
    run_label3.font.bold = True
    run_label3.font.color.rgb = RGBColor(*s["color"])
    
    p_value3 = cell_value3.paragraphs[0]
    p_value3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_text_with_highlight(p_value3, ausbildung, s["font"], s["size"], s["color"])
    
    # Merge image cells vertically across all 3 rows and add placeholder
    try:
        table.cell(0, 2).merge(table.cell(2, 2))
        merged_cell = table.cell(0, 2)
        merged_cell.text = ""
        p_merged = merged_cell.paragraphs[0]
        p_merged.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_merged = p_merged.add_run(get_text(translations, "cv", "insert_image", language))
        run_merged.font.name = s["font"]
        run_merged.font.size = Pt(s["size"])
        run_merged.font.color.rgb = RGBColor(*s["color"])
        from docx.enum.text import WD_COLOR_INDEX
        run_merged.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception:
        # Fallback: attempt element merge for older versions
        try:
            cell_image._element.merge(cell_image2._element)
            cell_image._element.merge(cell_image3._element)
        except Exception:
            pass


def add_fachwissen_table(doc, skills_data):
    """
    Render Fachwissen_und_Schwerpunkte as a 2-column table.
    Column 1: Category (bold) - 20%
    Column 2: Bullet list of Inhalt items - 80%
    """
    from docx.shared import Inches
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls
    
    if not skills_data:
        return
    
    # Create table: one row per category
    table = doc.add_table(rows=len(skills_data), cols=2)
    
    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    
    tblBorders = parse_xml(r'<w:tBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'</w:tBorders>' % nsdecls('w'))
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)
    
    # Set column widths: 20% for category, 80% for content
    available_width = get_available_width(doc)
    widths = [available_width * 0.20, available_width * 0.80]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    s_text = styles["text"]
    s_bullet = styles["bullet"]
    
    # Fill rows
    for row_idx, item in enumerate(skills_data):
        kategorie = item.get("Kategorie", "")
        inhalt = item.get("Inhalt", item.get("BulletList", []))  # Support both formats
        
        cell_cat = table.rows[row_idx].cells[0]
        cell_list = table.rows[row_idx].cells[1]
        
        remove_cell_borders(cell_cat)
        remove_cell_borders(cell_list)
        
        # Clear and populate category cell
        cell_cat.text = ""
        p_cat = cell_cat.paragraphs[0]
        p_cat.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p_cat.paragraph_format.space_before = Pt(0)
        p_cat.paragraph_format.space_after = Pt(0)
        run_cat = p_cat.add_run(kategorie)
        run_cat.font.name = s_text["font"]
        run_cat.font.size = Pt(s_text["size"])
        run_cat.font.bold = True
        run_cat.font.color.rgb = RGBColor(*s_text["color"])
        
        # Set vertical alignment to top for category cell
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        cell_cat.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        
        # Clear and populate content cell with comma-separated list
        cell_list.text = ""
        p = cell_list.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = s_text.get("line_spacing", 1.0)
        
        # Join all items with comma and space
        comma_separated = ", ".join(inhalt)
        
        # Add as single run
        run = p.add_run(comma_separated)
        run.font.name = s_text["font"]
        run.font.size = Pt(s_text["size"])
        run.font.color.rgb = RGBColor(*s_text["color"])


def add_education_table(doc, education_data):
    """
    Render Aus_und_Weiterbildung as a 2-column table.
    Column 1: Time Range (YYYY - YYYY or single year)
    Column 2: Row1 = Institution, Row2 = Ausbildung/Titel
    """
    from docx.shared import Inches
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    if not education_data:
        return
    
    # Create table: one row per education item
    table = doc.add_table(rows=len(education_data), cols=2)
    
    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    
    tblBorders = parse_xml(r'<w:tBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'</w:tBorders>' % nsdecls('w'))
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)
    
    # Set column widths: 20% for time range, 80% for institution/title
    available_width = get_available_width(doc)
    widths = [available_width * 0.20, available_width * 0.80]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    s_text = styles["text"]
    
    # Fill rows
    for row_idx, item in enumerate(education_data):
        zeitraum = item.get("Zeitraum", "")
        institution = item.get("Institution", "")
        ausbildung_titel = item.get("Abschluss", "")
        
        cell_time = table.rows[row_idx].cells[0]
        cell_info = table.rows[row_idx].cells[1]
        
        remove_cell_borders(cell_time)
        remove_cell_borders(cell_info)
        
        # Column 1: Time Range
        cell_time.text = ""
        p_time = cell_time.paragraphs[0]
        p_time.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        add_text_with_highlight(p_time, zeitraum, s_text["font"], s_text["size"], s_text["color"])
        
        # Column 2: Institution and Title in one paragraph separated by soft line break
        cell_info.text = ""
        p_info = cell_info.paragraphs[0]
        p_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Institution (bold)
        add_text_with_highlight(p_info, institution, s_text["font"], s_text["size"], s_text["color"], bold=True)

        # soft line break (shift+Enter)
        from docx.enum.text import WD_BREAK
        p_info.runs[-1].add_break(WD_BREAK.LINE)

        # Title on the next visual line (same paragraph)
        add_text_with_highlight(p_info, ausbildung_titel, s_text["font"], s_text["size"], s_text["color"])

        # Add spacing after the combined paragraph so rows are visually separated
        p_info.paragraph_format.space_after = Pt(6)


def add_trainings_table(doc, trainings_data):
    """
    Render Trainings_und_Zertifizierungen as a 2-column table.
    Column 1: Time Range (YYYY - YYYY)
    Column 2: Row1 = Institution (bold), Row2 = Ausbildung/Titel or Zertifizierung
    """
    from docx.shared import Inches
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    if not trainings_data:
        return

    table = doc.add_table(rows=len(trainings_data), cols=3)

    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)

    tblBorders = parse_xml(r'<w:tBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'</w:tBorders>' % nsdecls('w'))
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)

    # Column widths: 20% time, 20% institution, 60% certification/title
    available_width = get_available_width(doc)
    widths = [available_width * 0.20, available_width * 0.20, available_width * 0.60]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    s_text = styles["text"]

    for row_idx, item in enumerate(trainings_data):
        zeitraum = item.get("Zeitraum", "")
        institution = item.get("Institution", "")
        titel = item.get("Titel", item.get("Ausbildung_Titel", ""))

        cell_time = table.rows[row_idx].cells[0]
        cell_inst = table.rows[row_idx].cells[1]
        cell_title = table.rows[row_idx].cells[2]

        remove_cell_borders(cell_time)
        remove_cell_borders(cell_inst)
        remove_cell_borders(cell_title)

        # Time
        cell_time.text = ""
        p_time = cell_time.paragraphs[0]
        add_text_with_highlight(p_time, zeitraum, s_text["font"], s_text["size"], s_text["color"])
        p_time.paragraph_format.space_after = Pt(3)  # Reduced spacing

        # Institution (bold)
        cell_inst.text = ""
        p_inst = cell_inst.paragraphs[0]
        add_text_with_highlight(p_inst, institution, s_text["font"], s_text["size"], s_text["color"], bold=True)
        p_inst.paragraph_format.space_after = Pt(3)  # Reduced spacing

        # Certification / Training (title)
        cell_title.text = ""
        p_t = cell_title.paragraphs[0]
        add_text_with_highlight(p_t, titel, s_text["font"], s_text["size"], s_text["color"])
        p_t.paragraph_format.space_after = Pt(3)  # Reduced spacing


# -------------------------------------
# Hauptfunktion zum Erstellen des CVs
# -------------------------------------

def add_header_with_logo(doc):
    """Add a header with 3-column table: logo (40%), empty (10%), text (50%)."""
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    import os
    
    header_config = styles.get("header", {})
    logo_path = header_config.get("logo_path", "")
    logo_width = header_config.get("logo_width_cm", 4.0)
    logo_height = header_config.get("logo_height_cm", 1.5)
    header_text = header_config.get("text", "")
    text_font = header_config.get("text_font", "Aptos")
    text_size = header_config.get("text_size", 10)
    text_color = header_config.get("text_color", [68, 68, 68])
    
    # Access the header section
    section = doc.sections[0]
    header = section.header
    
    # Use the first paragraph for the table
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.text = ""
    
    # Create table in document (will be moved to header manually)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Set column widths: 40%, 10%, 50%
    available_width = get_available_width(doc)
    col_widths = [available_width * 0.40, available_width * 0.10, available_width * 0.50]
    for idx, width in enumerate(col_widths):
        table.columns[idx].width = int(width)
    
    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    
    tblBorders = parse_xml(r'<w:tBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'</w:tBorders>' % nsdecls('w'))
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)
    
    for cell in table.rows[0].cells:
        remove_cell_borders(cell)
    
    # Column 1: Logo (left aligned)
    cell_logo = table.rows[0].cells[0]
    cell_logo.text = ""
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if logo_path:
        full_logo_path = abs_path(logo_path)
        if os.path.exists(full_logo_path):
            try:
                run = p_logo.add_run()
                # Only set width, let Word calculate height to maintain aspect ratio
                run.add_picture(full_logo_path, width=Cm(logo_width))
                print(f"✅ Logo erfolgreich eingefügt: {full_logo_path}")
            except Exception as e:
                print(f"⚠️  Logo konnte nicht geladen werden: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"⚠️  Logo-Datei nicht gefunden: {full_logo_path}")
    
    # Column 2: Empty (10%)
    cell_empty = table.rows[0].cells[1]
    cell_empty.text = ""
    
    # Column 3: Text (right aligned)
    cell_text = table.rows[0].cells[2]
    cell_text.text = ""
    p_text = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if header_text:
        run_text = p_text.add_run(header_text)
        run_text.font.name = text_font
        run_text.font.size = Pt(text_size)
        run_text.font.color.rgb = RGBColor(*text_color)
    
    # Move table to header
    table_element = table._element
    # Find the table in the body and get its parent
    body_element = doc._body._element
    if table_element in body_element:
        body_element.remove(table_element)
    # Insert into header
    header._element.insert(0, table_element)
    
    # Remove default empty paragraph in header if it exists
    if header.paragraphs and len(header.paragraphs) > 0:
        if not header.paragraphs[0].text.strip():
            p_element = header.paragraphs[0]._element
            if p_element.getparent() is not None:
                p_element.getparent().remove(p_element)


def generate_cv(json_path, output_dir=None, interactive=True, language="de"):

    json_path = abs_path(json_path)
    translations = load_translations()

    # JSON einlesen
    with open(json_path, 'r', encoding="utf-8") as f:
        data = json.load(f)
    
    # JSON-Struktur validieren
    critical, info = validate_json_structure(data, language=language)
    if (critical or info) and interactive:
        # Build warning message with explanation
        warning_msg = get_text(translations, 'validation', 'warning_msg', language)
        
        # Build details with better formatting and icons
        details_parts = []
        
        if critical:
            details_parts.append(get_text(translations, 'validation', 'critical_problems', language))
            details_parts.append("─" * 40)
            for err in critical:
                # Add specific icons based on error type
                if get_text(translations, 'validation', 'missing_field', language) in err:
                    icon = "❌"
                elif get_text(translations, 'validation', 'must_be_array', language) in err or get_text(translations, 'validation', 'must_be_object', language) in err:
                    icon = "⚠️"
                elif "Ungültig" in err or "ungültig" in err or "Invalid" in err or "non valide" in err:
                    icon = "🚫"
                else:
                    icon = "🔴"
                details_parts.append(f"  {icon} {err}")
            
        if info:
            if critical:
                details_parts.append("")  # Empty line separator
            details_parts.append(get_text(translations, 'validation', 'info_problems', language))
            details_parts.append("─" * 40)
            for wrn in info:
                # Add specific icons based on warning type
                if "sollte" in wrn or "should" in wrn or "devrait" in wrn:
                    icon = "💡"
                elif "Typ" in wrn or "sein" in wrn or "string" in wrn or "chaîne" in wrn:
                    icon = "ℹ️"
                else:
                    icon = "🟡"
                details_parts.append(f"  {icon} {wrn}")
        
        details = "\n".join(details_parts)
        
        try:
            proceed = show_warning(warning_msg, title=get_text(translations, 'validation', 'warning_title', language), details=details)
            if not proceed:
                print(f"❌ {get_text(translations, 'validation', 'user_cancelled', language)}")
                return None
        except Exception as e:
            print(f"Warnung konnte nicht angezeigt werden: {e}")
            print("Probleme gefunden:", "\n".join(critical + info))
            user_input = input(get_text(translations, 'validation', 'proceed_anyway', language))
            if user_input.lower() not in ['j', 'ja', 'y', 'yes', 'o', 'oui']:
                return None

    # Lade Template-Datei mit Header/Footer oder erstelle leeres Dokument
    template_path = abs_path("../templates/cv_template.docx")
    if os.path.exists(template_path):
        doc = Document(template_path)
        # Template bringt bereits Header, Footer und Seitenränder mit
    else:
        print(f"Warnung: {get_text(translations, 'validation', 'no_template', language)}")
        doc = Document()
        # Set page margins: 1.5 cm all sides and A4 size
        from docx.shared import Cm
        sections = doc.sections
        for section in sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        # Add header with logo
        add_header_with_logo(doc)

    firstname = str(data.get("Vorname", "Unbekannt"))
    lastname = str(data.get("Nachname", "Unbekannt"))

    # Remove default empty paragraph if it exists
    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p_element = doc.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    # -----------------------------
    # Überschrift 1 (Name)
    # -----------------------------
    add_heading_1(doc, f"{firstname} {lastname}")

    # -----------------------------
    # Basisinfos
    # -----------------------------
    hauptrolle = data.get("Hauptrolle", {})
    rolle_desc = hauptrolle.get("Beschreibung", "") if isinstance(hauptrolle, dict) else str(hauptrolle)
    add_basic_info_table(doc, 
                         rolle_desc,
                         str(data.get("Nationalität", "")),
                         str(data.get("Ausbildung", "")),
                         language=language,
                         translations=translations)

    # -----------------------------
    # Kurzprofil
    # -----------------------------
    add_heading_1(doc, get_text(translations, 'cv', 'profile', language))
    add_normal_text(doc, str(data.get("Kurzprofil", "")))

    # -----------------------------
    # Expertise & Fachwissen
    # -----------------------------
    add_heading_1(doc, get_text(translations, 'cv', 'expertise', language))
    add_heading_2(doc, get_text(translations, 'cv', 'skills', language))
    skills = data.get("Fachwissen_und_Schwerpunkte", [])
    add_fachwissen_table(doc, skills)

    # -----------------------------
    # Aus- und Weiterbildung
    # -----------------------------
    add_heading_2(doc, get_text(translations, 'cv', 'education', language))
    education = data.get("Aus_und_Weiterbildung", [])
    add_education_table(doc, education)

    # -----------------------------
    # Trainings & Zertifizierungen
    # -----------------------------
    add_heading_2(doc, get_text(translations, 'cv', 'certifications', language))
    trainings = data.get("Trainings_und_Zertifizierungen", [])
    add_trainings_table(doc, trainings)

    # -----------------------------
    # Sprachen
    # -----------------------------
    add_heading_2(doc, get_text(translations, 'cv', 'languages', language))
    sprachen = data.get("Sprachen", [])
    add_sprachen_table(doc, sprachen, language=language, translations=translations)

    # -----------------------------
    # Referenzprojekte
    # -----------------------------
    doc.add_page_break()
    add_heading_1(doc, get_text(translations, 'cv', 'projects', language))
    referenzprojekte = data.get("Ausgewählte_Referenzprojekte", [])
    for projekt in referenzprojekte:
        add_referenzprojekt_section(doc, projekt, language=language, translations=translations)


    # -----------------------------
    # Dateien speichern
    # -----------------------------
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if output_dir is None:
        # Fallback to old structure if not provided
        out_docx = abs_path(f"../output/word/cv_{firstname}_{lastname}_{timestamp}.docx")
    else:
        out_docx = os.path.join(output_dir, f"cv_{firstname}_{lastname}_{timestamp}.docx")

    # Vor dem Speichern: Highlight alle fehlenden Daten im gesamten Dokument
    highlight_missing_data_in_document(doc, translations=translations, language=language)
    
    doc.save(out_docx)
    print(f"✅ Word-Datei erstellt: {out_docx}")
    return out_docx


# Old dialog removed - now using modern_dialogs.select_json_file() directly


def select_json_file():
    """Oeffnet einen Datei-Dialog zur Auswahl einer JSON-Datei (standardmässig im output/-Ordner)"""
    try:
        from dialogs import select_json_file as picker
    except ImportError:
        from scripts.dialogs import select_json_file as picker
    # Default to output/ for manual selection
    output_dir = abs_path("../output/")
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    return picker("Wählen Sie eine JSON-Datei für den CV", initialdir=output_dir)


def add_sprachen_table(doc, sprachen_data, language="de", translations=None):
    """
    Render Sprachen as a 6-column table (2 rows max).
    Columns: Sprache | Level (stars) | Sprache | Level (stars) | Sprache | Level (stars)
    Sorted by star count descending, max 6 languages (2 rows × 3 pairs)
    Stars: ★ = full, ☆ = empty
    """
    from docx.shared import Inches
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls
    
    if not sprachen_data:
        return
    
    # Define level to star count mapping (supports both strings and numbers)
    # We include translations of the levels to ensure mapping works regardless of target language
    level_stars = {
        # German (Original)
        "Muttersprache": 5,
        "Verhandlungssicher": 4,
        "Sehr gute Kenntnisse": 3,
        "Gute Kenntnisse": 2,
        "Grundkenntnisse": 1,
        # Translations from translations.json
        get_text(translations, 'levels', 'level_5', language): 5,
        get_text(translations, 'levels', 'level_4', language): 4,
        get_text(translations, 'levels', 'level_3', language): 3,
        get_text(translations, 'levels', 'level_2', language): 2,
        get_text(translations, 'levels', 'level_1', language): 1,
        # English variants often found
        "Native speaker": 5,
        "Native": 5,
        "Full professional proficiency": 4,
        "Professional proficiency": 4,
        "Very good knowledge": 3,
        "Good knowledge": 2,
        "Basic knowledge": 1,
        # Numeric
        5: 5,
        4: 4,
        3: 3,
        2: 2,
        1: 1
    }
    
    # Sort languages by star count descending
    sorted_sprachen = sorted(sprachen_data, 
                           key=lambda x: level_stars.get(x.get("Level", ""), 0), 
                           reverse=True)
    
    # Take max 6 languages (2 rows × 3 pairs)
    display_sprachen = sorted_sprachen[:6]
    
    # Create table: 2 rows, 6 columns
    table = doc.add_table(rows=2, cols=6)
    
    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    
    tblBorders = parse_xml(r'<w:tBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'</w:tBorders>' % nsdecls('w'))
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)
    
    # Set column widths: equal distribution
    available_width = get_available_width(doc)
    col_width = available_width / 6
    for row in table.rows:
        for idx in range(6):
            cell = row.cells[idx]
            cell.width = col_width
            set_cell_padding_zero(cell)
        # Set smaller row height for compact display
        from docx.shared import Pt
        row.height = Pt(18)  # Smaller row height
        row.height_rule = 2  # WD_ROW_HEIGHT_RULE.AT_LEAST
    
    # Fill table cells
    s_text = styles["text"]
    
    for i, sprache in enumerate(display_sprachen):
        row_idx = i // 3  # 0 or 1
        col_start = (i % 3) * 2  # 0, 2, 4
        
        if row_idx >= 2:  # Safety check
            break
            
        # Sprache cell
        cell_sprache = table.rows[row_idx].cells[col_start]
        cell_sprache.text = ""
        p_sprache = cell_sprache.paragraphs[0]
        p_sprache.paragraph_format.line_spacing = 0.8  # Tighter spacing
        run_sprache = p_sprache.add_run(sprache.get("Sprache", ""))
        run_sprache.font.name = s_text["font"]
        run_sprache.font.size = Pt(s_text["size"])
        run_sprache.font.color.rgb = RGBColor(*s_text["color"])
        
        # Level cell (stars)
        cell_level = table.rows[row_idx].cells[col_start + 1]
        cell_level.text = ""
        p_level = cell_level.paragraphs[0]
        p_level.paragraph_format.line_spacing = 0.8  # Tighter spacing
        
        level = sprache.get("Level", "")
        star_count = parse_level(level)
        
        # Add full stars
        for _ in range(star_count):
            run_star = p_level.add_run("★")
            run_star.font.name = s_text["font"]
            run_star.font.size = Pt(s_text["size"])
            run_star.font.color.rgb = RGBColor(255, 121, 0)  # Orange color
        
        # Add empty stars
        for _ in range(5 - star_count):
            run_empty = p_level.add_run("☆")
            run_empty.font.name = s_text["font"]
            run_empty.font.size = Pt(s_text["size"])
            run_empty.font.color.rgb = RGBColor(128, 128, 128)  # Gray color

            run_empty.font.color.rgb = RGBColor(128, 128, 128)  # Gray color


def add_referenzprojekt_section(doc, projekt, language="de", translations=None):
    """
    Render a single reference project as a unified 5-row table block:
    Row 1: Kunde (100% width, Heading 2)
    Row 2: Zeitraum (20%) | Rolle (80%) (text style)
    Row 3: Tätigkeiten (100% width, bullets)
    Row 4: Technologien Titel (20%) | Technologien Inhalt (80%) (text style)
    Row 5: Methodik Titel (20%) | Methodik Inhalt (80%) (text style)
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Pt, Cm, RGBColor
    
    kunde = projekt.get("Kunde", "")
    zeitraum = projekt.get("Zeitraum", "")
    rolle = projekt.get("Rolle", "")
    taetigkeiten = projekt.get("Tätigkeiten", [])
    technologien = projekt.get("Technologien", "")
    methodik = projekt.get("Methodik", "")
    
    # Create unified 5-row table with 2 columns
    table = doc.add_table(rows=5, cols=2)
    
    # Remove table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr %s></w:tblPr>' % nsdecls('w'))
        tbl.insert(0, tblPr)
    
    tblBorders = parse_xml(r'<w:tBorders %s>'
                          r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                          r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'  
                          r'</w:tBorders>' % nsdecls('w'))
    existing_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tBorders')
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(tblBorders)
    
    # Set column widths: 20% for col1, 80% for col2
    available_width = get_available_width(doc)
    widths = [available_width * 0.20, available_width * 0.80]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    
    # Row 1: Kunde (merge both columns for 100% width)
    row1_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    remove_cell_borders(row1_cell)
    row1_cell.text = ""
    p_kunde = row1_cell.paragraphs[0]
    s_heading2 = styles["heading2"]
    add_text_with_highlight(p_kunde, kunde, s_heading2["font"], s_heading2["size"], s_heading2["color"], bold=s_heading2.get("bold", False))
    p_kunde.paragraph_format.space_before = Pt(0)
    p_kunde.paragraph_format.space_after = Pt(3)
    
    # Row 2: Zeitraum (20%) | Rolle (80%)
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    row2 = table.rows[1]
    row2.height = Cm(0.7)
    row2.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    
    cell_zeitraum = row2.cells[0]
    cell_rolle = row2.cells[1]
    remove_cell_borders(cell_zeitraum)
    remove_cell_borders(cell_rolle)
    
    s_heading2 = styles["heading2"]
    s_text = styles["text"]
    
    cell_zeitraum.text = ""
    p_zeitraum = cell_zeitraum.paragraphs[0]
    add_text_with_highlight(p_zeitraum, zeitraum, s_heading2["font"], s_heading2["size"], s_heading2["color"], bold=s_heading2.get("bold", False))
    p_zeitraum.paragraph_format.space_before = Pt(0)
    p_zeitraum.paragraph_format.space_after = Pt(0)
    
    cell_rolle.text = ""
    p_rolle = cell_rolle.paragraphs[0]
    add_text_with_highlight(p_rolle, rolle, s_heading2["font"], s_heading2["size"], s_heading2["color"], bold=s_heading2.get("bold", False))
    p_rolle.paragraph_format.space_before = Pt(0)
    p_rolle.paragraph_format.space_after = Pt(0)
    
    # Row 3: Tätigkeiten (merge both columns for 100% width)
    row3_cell = table.rows[2].cells[0].merge(table.rows[2].cells[1])
    remove_cell_borders(row3_cell)
    row3_cell.text = ""
    
    # Add activities as bullets in the merged cell
    if taetigkeiten:
        # Clear default paragraph
        row3_cell._element.clear_content()
        
        s_bullet = styles["bullet"]
        
        for taetigkeit in taetigkeiten:
            if taetigkeit.strip():
                p = row3_cell.add_paragraph()
                
                # Set up hanging indent for proper text alignment
                # Left indent for all lines
                left_indent_value = Pt(18)  # Position where text starts
                first_line_negative = Pt(-18)  # Pull first line back to position bullet
                
                p.paragraph_format.left_indent = left_indent_value
                p.paragraph_format.first_line_indent = first_line_negative
                
                # Add bullet symbol from styles.json
                bullet_symbol = s_bullet.get("symbol", "■")
                symbol_run = p.add_run(bullet_symbol)
                symbol_run.font.name = s_bullet["font"]
                symbol_run.font.size = Pt(s_bullet.get("symbol_size", s_bullet["size"]))
                symbol_run.font.color.rgb = RGBColor(*s_bullet["color"])
                
                # Add tab to separate symbol from text
                p.add_run("\t")
                
                # Add text with normal size
                add_text_with_highlight(p, taetigkeit, s_text["font"], s_text["size"], s_text["color"])
                
                p.paragraph_format.space_before = Pt(s_bullet.get("space_before", 0))
                p.paragraph_format.space_after = Pt(s_bullet.get("space_after", 3))
                p.paragraph_format.line_spacing = s_bullet.get("line_spacing", 1.0)
                
                # Set tab stop at the text start position
                from docx.shared import Pt
                from docx.enum.text import WD_TAB_ALIGNMENT
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(left_indent_value, WD_TAB_ALIGNMENT.LEFT)
    
    # Row 4: Technologien (label 20% | content 80%)
    cell_tech_label = table.rows[3].cells[0]
    cell_tech_content = table.rows[3].cells[1]
    remove_cell_borders(cell_tech_label)
    remove_cell_borders(cell_tech_content)
    
    cell_tech_label.text = ""
    p_tech_label = cell_tech_label.paragraphs[0]
    run_tech_label = p_tech_label.add_run(get_text(translations, 'cv', 'tech_label', language))
    run_tech_label.font.name = s_text["font"]
    run_tech_label.font.size = Pt(s_text["size"])
    run_tech_label.font.color.rgb = RGBColor(*s_text["color"])
    run_tech_label.italic = True
    run_tech_label.bold = False
    p_tech_label.paragraph_format.space_before = Pt(0)
    p_tech_label.paragraph_format.space_after = Pt(0)
    
    cell_tech_content.text = ""
    p_tech_content = cell_tech_content.paragraphs[0]
    add_text_with_highlight(p_tech_content, technologien, s_text["font"], s_text["size"], s_text["color"], italic=True)
    p_tech_content.paragraph_format.space_before = Pt(0)
    p_tech_content.paragraph_format.space_after = Pt(0)
    
    # Row 5: Methodik (label 20% | content 80%)
    cell_meth_label = table.rows[4].cells[0]
    cell_meth_content = table.rows[4].cells[1]
    remove_cell_borders(cell_meth_label)
    remove_cell_borders(cell_meth_content)
    
    cell_meth_label.text = ""
    p_meth_label = cell_meth_label.paragraphs[0]
    run_meth_label = p_meth_label.add_run(get_text(translations, 'cv', 'methodology_label', language))
    run_meth_label.font.name = s_text["font"]
    run_meth_label.font.size = Pt(s_text["size"])
    run_meth_label.font.color.rgb = RGBColor(*s_text["color"])
    run_meth_label.italic = True
    run_meth_label.bold = False
    p_meth_label.paragraph_format.space_before = Pt(0)
    p_meth_label.paragraph_format.space_after = Pt(0)
    
    cell_meth_content.text = ""
    p_meth_content = cell_meth_content.paragraphs[0]
    add_text_with_highlight(p_meth_content, methodik, s_text["font"], s_text["size"], s_text["color"], italic=True)
    p_meth_content.paragraph_format.space_before = Pt(0)
    p_meth_content.paragraph_format.space_after = Pt(0)
    
    # Add spacing after the project block by adding an empty paragraph
    spacing_paragraph = doc.add_paragraph()
    spacing_paragraph.paragraph_format.space_after = Pt(24)


# Old dialog removed - now using modern_dialogs functions


# Lade Stile global nach allen Definitionen
styles = load_styles("styles.json")
if __name__ == "__main__":
    # Import dialogs - handle both direct execution and module import
    try:
        from dialogs import show_success, ask_yes_no, ModernDialog
    except ImportError:
        from scripts.dialogs import show_success, ask_yes_no, ModernDialog
    
    json_file = select_json_file()
    if json_file:
        output_file = generate_cv(json_file)
        if output_file:
            # Show success with open button
            output_dir = os.path.dirname(output_file)
            filename = os.path.basename(output_file)
            
            details = (
                f"{ModernDialog.ICON_WORD} Output Directory:\n"
                f"  {output_dir}\n\n"
                f"Filename:\n"
                f"  {filename}"
            )
            
            result = show_success(
                "Das CV-Dokument wurde erfolgreich erstellt.",
                title="CV erfolgreich generiert",
                details=details,
                file_path=output_file
            )
            
            # Open document if user clicked "Open"
            if result == 'open':
                try:
                    import platform
                    import subprocess
                    
                    if platform.system() == 'Windows':
                        os.startfile(output_file)
                    elif platform.system() == 'Darwin':  # macOS
                        subprocess.run(['open', output_file])
                    else:  # Linux
                        subprocess.run(['xdg-open', output_file])
                except Exception as e:
                    print(f"Fehler beim Öffnen der Datei: {e}")
        else:
            print("❌ JSON-Validierung fehlgeschlagen. Programm abgebrochen.")
    else:
        print("❌ Programm abgebrochen.")
