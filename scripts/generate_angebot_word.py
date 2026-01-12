import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta
import calendar

def hex_to_rgb(hex_color):
    """Converts hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell"""
    cell_properties = cell._element.get_or_add_tcPr()
    shading_element = OxmlElement('w:shd')
    shading_element.set(qn('w:val'), 'clear')
    shading_element.set(qn('w:color'), 'auto')
    shading_element.set(qn('w:fill'), color_hex)
    cell_properties.append(shading_element)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def set_cell_padding(cell, padding_dxa=28):
    """Sets all internal cell margins/padding (default 28 dxa = 0.05 cm)"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(padding_dxa))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_cell_borders(cell):
    """Entfernt alle Rahmen einer Tabellenzelle via XML"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def get_default_validity_date():
    """Calculates default validity: end of month, or +10 days if <10 days left"""
    today = datetime.now()
    _, last_day = calendar.monthrange(today.year, today.month)
    end_of_month = today.replace(day=last_day)
    
    if (end_of_month - today).days < 10:
        return today + timedelta(days=10)
    return end_of_month

def abs_path(relative_path):
    """Returns absolute path relative to this script's directory"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(script_dir, relative_path)

def load_styles():
    """Loads styles from styles.json"""
    styles_path = abs_path("styles.json")
    if os.path.exists(styles_path):
        with open(styles_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

def add_paragraph_with_bold(parent_obj, text, style=None):
    """
    Adds a paragraph to a document or cell, parsing **text** as bold.
    """
    if hasattr(parent_obj, 'add_paragraph'):
        p = parent_obj.add_paragraph(style=style)
    else:
        p = parent_obj # Assume it is already a paragraph
        
    parts = text.split('**')
    for i, part in enumerate(parts):
        chunk = part
        is_bold = (i % 2 == 1)
        
        # Check for highlighting if "bitte prüfen" is present
        highlight_needed = "bitte prüfen" in chunk.lower()
        
        run = p.add_run(chunk)
        if is_bold:
            run.bold = True
        if highlight_needed:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    return p

def add_bullet_paragraph(doc, text):
    """Adds a custom bullet paragraph using styles.json config"""
    styles = load_styles()
    s_bullet = styles.get("bullet", {})
    text_cfg = styles.get("text", {})
    
    p = doc.add_paragraph()
    indent = s_bullet.get("indent", 0)
    hanging = s_bullet.get("hanging_indent", 12)
    p.paragraph_format.left_indent = Pt(indent + hanging)
    p.paragraph_format.first_line_indent = Pt(-hanging)
    p.paragraph_format.space_before = Pt(s_bullet.get("space_before", 0))
    p.paragraph_format.space_after = Pt(s_bullet.get("space_after", 3))
    p.paragraph_format.line_spacing = s_bullet.get("line_spacing", 1.0)
    
    # Add tab stop for bullet alignment
    p.paragraph_format.tab_stops.add_tab_stop(Pt(indent + hanging))

    # 1. Add Bullet Symbol
    symbol = s_bullet.get("symbol", "■")
    run_sym = p.add_run(f"{symbol}\t")
    run_sym.font.name = s_bullet.get("font", "Aptos")
    run_sym.font.size = Pt(s_bullet.get("symbol_size", 9))
    if "color" in s_bullet:
        run_sym.font.color.rgb = RGBColor(*s_bullet["color"])

    # 2. Add content with bold support
    parts = text.split('**')
    for i, part in enumerate(parts):
        chunk = part
        is_bold = (i % 2 == 1)
        
        run = p.add_run(chunk)
        run.font.name = text_cfg.get("font", "Aptos")
        run.font.size = Pt(text_cfg.get("size", 11))
        if is_bold:
            run.bold = True
        
        # Highlighting for check markers
        if "bitte prüfen" in chunk.lower():
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    return p

def add_criteria_table(doc, title, criteria_list):
    """Refactored helper for adding a criteria table to word"""
    if not criteria_list:
        return
    
    # Load styles
    styles_cfg = load_styles()
    text_cfg = styles_cfg.get("text", {})
    table_cfg = styles_cfg.get("table", {})
    width_total = table_cfg.get("width_cm", 17.0)
    ratios = table_cfg.get("column_ratios_criteria", [0.4, 0.2, 0.4])
    bg_color = table_cfg.get("header_bg_color", "EEEEEE")
    f_size = table_cfg.get("font_size", 10)
    f_name = text_cfg.get("font", "Aptos")
        
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=3)
    table.style = table_cfg.get("border_style", "Table Grid")
    
    # Enable autofit off for precise control
    table.autofit = False
    table.width = Cm(width_total)
    
    # Calculate widths based on ratios
    widths = [Cm(width_total * r) for r in ratios]
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kriterium'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Hinweis'
    
    # Set widths for header cells
    for i, width in enumerate(widths):
        hdr_cells[i].width = width

    # Header styling & Font size
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, bg_color)
        set_cell_padding(cell)
        p = cell.paragraphs[0]
        if i == 1: # Status column
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        run.font.name = f_name
        run.font.bold = True
        run.font.size = Pt(f_size)

    for item in criteria_list:
        row_cells = table.add_row().cells
        
        # Explicitly set cell widths for each new row
        for i, width in enumerate(widths):
            row_cells[i].width = width
            set_cell_padding(row_cells[i])
            
        # 1. Kriterium
        row_cells[0].text = item.get("kriterium", "")
        run_crit = row_cells[0].paragraphs[0].runs[0]
        run_crit.font.name = f_name
        run_crit.font.size = Pt(f_size)
        
        status = str(item.get("erfuellt", "")).lower().strip()
        
        # Mapping to display text and color (with icons)
        status_map = {
            "erfüllt": ("✅ Erfüllt", RGBColor(39, 174, 96), False),
            "teilweise erfüllt": ("⚠️ Teilweise", RGBColor(243, 156, 18), True),
            "potenziell erfüllt": ("⚠️ Teilweise", RGBColor(243, 156, 18), True),
            "nicht erfüllt": ("❌ Nicht erfüllt", RGBColor(192, 57, 43), True),
            "nicht explizit erwähnt": ("⚪ Nicht erwähnt", RGBColor(127, 140, 141), False),
            "true": ("✅ Erfüllt", RGBColor(39, 174, 96), False),
            "false": ("❌ Nicht erfüllt", RGBColor(192, 57, 43), True),
            "! bitte prüfen !": ("❓ Prüfen", RGBColor(0, 0, 0), True)
        }
        
        display_text, color, is_bold = status_map.get(status, (status.capitalize(), RGBColor(0, 0, 0), False))
        
        # 2. Status
        row_cells[1].text = display_text
        p_status = row_cells[1].paragraphs[0]
        p_status.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_status = p_status.runs[0]
        run_status.font.name = f_name
        run_status.font.color.rgb = color
        run_status.font.bold = is_bold
        run_status.font.size = Pt(f_size)
            
        # 3. Hinweis
        row_cells[2].text = item.get("begruendung", "")
        run_hint = row_cells[2].paragraphs[0].runs[0]
        run_hint.font.name = f_name
        run_hint.font.size = Pt(f_size)

def generate_angebot_word(json_path, output_path):
    """
    Generates a Word document for the Offer based on the JSON data.
    """
    # Load Data
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Load External Styles
    styles_cfg = load_styles()
    
    def get_color(cfg_section, key="color"):
        rgb_list = cfg_section.get(key, [0, 0, 0])
        return RGBColor(*rgb_list)

    doc = Document()
    
    # --- Page Setup (A4 & Margins 2cm) ---
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # --- Styles Setup ---
    # Normal Text (Normal)
    style_cfg = styles_cfg.get("text", {})
    style = doc.styles['Normal']
    font = style.font
    font.name = style_cfg.get("font", "Aptos")
    font.size = Pt(style_cfg.get("size", 11))
    font.color.rgb = get_color(style_cfg)
    style.paragraph_format.space_before = Pt(style_cfg.get("space_before", 0))
    style.paragraph_format.space_after = Pt(style_cfg.get("space_after", 6))

    # Heading 1 (Titel)
    h1_cfg = styles_cfg.get("heading1", {})
    h1 = doc.styles['Heading 1']
    h1.font.name = h1_cfg.get("font", "Aptos")
    h1.font.size = Pt(h1_cfg.get("size", 16))
    h1.font.bold = h1_cfg.get("bold", True)
    h1.font.color.rgb = get_color(h1_cfg)
    h1.paragraph_format.space_before = Pt(h1_cfg.get("space_before", 5))
    h1.paragraph_format.space_after = Pt(h1_cfg.get("space_after", 5))
    
    # Heading 2 (Abschnitte)
    h2_cfg = styles_cfg.get("heading2", {})
    h2 = doc.styles['Heading 2']
    h2.font.name = h2_cfg.get("font", "Aptos")
    h2.font.size = Pt(h2_cfg.get("size", 11))
    h2.font.bold = h2_cfg.get("bold", True)
    h2.font.color.rgb = get_color(h2_cfg)
    h2.paragraph_format.space_before = Pt(h2_cfg.get("space_before", 5))
    h2.paragraph_format.space_after = Pt(h2_cfg.get("space_after", 5))

    # Heading 3 (Unterabschnitte)
    try:
        h3 = doc.styles['Heading 3']
    except:
        h3 = doc.styles.add_style('Heading 3', 1)
    
    h3.font.name = h2_cfg.get("font", "Aptos")
    h3.font.size = Pt(h2_cfg.get("size", 11))
    h3.font.bold = True
    h3.font.color.rgb = get_color(h2_cfg) # Heading 2 color as requested for level 2 sections
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0) # Black for h3

    # --- Header ---
    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(1, 2, Cm(17.0))
    header_table.autofit = False
    
    # Left column for Logo, right for text
    header_table.rows[0].cells[0].width = Cm(10.0)
    header_table.rows[0].cells[1].width = Cm(7.0)
    
    # Logo (Placeholder logic - refers to styles.json if possible)
    header_cfg = styles_cfg.get("header", {})
    logo_filename = os.path.basename(header_cfg.get("logo_path", "logo.png"))
    logo_path = abs_path(f"../templates/{logo_filename}")
    
    cell_logo = header_table.cell(0, 0)
    if os.path.exists(logo_path):
        paragraph = cell_logo.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Cm(header_cfg.get("logo_width_cm", 4.0)))
    else:
        cell_logo.text = "Orange Business"

    cell_text = header_table.cell(0, 1)
    paragraph = cell_text.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run(header_cfg.get("text", "digital.orange-business.com"))
    run.font.color.rgb = get_color(header_cfg, "text_color")
    run.font.size = Pt(header_cfg.get("text_size", 10))
    run.font.name = header_cfg.get("text_font", "Aptos")

    # --- Footer ---
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = footer_paragraph.add_run()
    add_page_number(run)

    # --- Content ---
    
    meta = data.get("angebots_metadata", {})
    stellenbezug = data.get("stellenbezug", {})
    ansprechpartner = meta.get("ansprechpartner", {})
    
    # Placeholder logic
    contact_name = ansprechpartner.get("name", "")
    if not contact_name or contact_name.lower() in ["vorname nachname", "! bitte prüfen !"]:
        contact_name = "<Vor- Nachname>"
    
    contact_email = ansprechpartner.get("kontakt", "")
    if not contact_email or "@" not in contact_email:
        contact_email = "<E-Mail>"

    # 1. Letter Header (Sender & Recipient)
    # We use a 2-column table with no borders for the address layout
    letter_table = doc.add_table(rows=1, cols=2)
    letter_table.autofit = False
    letter_table.width = Cm(17.0)
    
    # Set column widths to 50/50
    for cell in letter_table.columns[0].cells:
        cell.width = Cm(8.5)
        set_cell_padding(cell)
    for cell in letter_table.columns[1].cells:
        cell.width = Cm(8.5)
        set_cell_padding(cell)
    
    # Left Cell: Sender & Date
    cell_sender = letter_table.cell(0, 0)
    p_sender = cell_sender.paragraphs[0]
    p_sender.add_run("Orange Business Digital").bold = True
    p_sender.add_run("\nIttigen, ").bold = False
    run_date = p_sender.add_run("DD.MM.YYYY")
    run_date.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    p_sender.add_run("\n\n\n\nIhr direkter Ansprechpartner:\n").bold = True
    run_name = p_sender.add_run(contact_name)
    if "<" in contact_name or "bitte prüfen" in contact_name.lower():
        run_name.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    p_sender.add_run("\n")
    run_mail = p_sender.add_run(contact_email)
    if "<" in contact_email or "bitte prüfen" in contact_email.lower():
        run_mail.font.highlight_color = WD_COLOR_INDEX.YELLOW
        
    p_sender.add_run("\n")
    run_mobile = p_sender.add_run("<Direkte Tel Nr.>")
    if "bitte prüfen" in run_mobile.text.lower() or "<" in run_mobile.text:
        run_mobile.font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    # Right Cell: Recipient
    cell_recipient = letter_table.cell(0, 1)
    cell_recipient.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_rec = cell_recipient.paragraphs[0]
    p_rec.add_run(meta.get("kunde", "")).bold = True
    run_rec_hint = p_rec.add_run("\n<bitte falls bekannt Ansprechpartner Kunde ergänzen>")
    run_rec_hint.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer

    # 2. Subject (Betreff)
    offer_id = meta.get("angebots_id", "GDJOB-<ID>")
    role_text = stellenbezug.get("rollenbezeichnung", "IT-Spezialist")
    org_text = stellenbezug.get("organisationseinheit", "")
    
    subject_p = doc.add_paragraph()
    
    # Header: Offer ID | Role
    if org_text:
        full_role = f"{role_text} {org_text}"
    else:
        full_role = role_text
        
    run_subj = subject_p.add_run(f"Angebot {offer_id} | {full_role}")
    run_subj.bold = True
    run_subj.font.size = Pt(12)
    if "<" in offer_id or "bitte prüfen" in offer_id.lower():
        run_subj.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_paragraph() # Spacer

    # Salutation
    salutation_p = doc.add_paragraph()
    salutation_p.add_run("Sehr geehrter ").bold = False
    run_salt = salutation_p.add_run("<Nachname Ansprechperson - bitte prüfen>")
    run_salt.font.highlight_color = WD_COLOR_INDEX.YELLOW
    salutation_p.paragraph_format.space_after = Pt(0)

    # 3. Ausgangslage & Kandidatenvorschlag (Direkt nach Briefanrede ohne Header)
    p_intro = add_paragraph_with_bold(doc, stellenbezug.get("kurzkontext", ""))
    p_intro.paragraph_format.space_before = Pt(0)
    
    kandidat = data.get("kandidatenvorschlag", {})
    add_paragraph_with_bold(doc, kandidat.get("eignungs_summary", ""))

    # 4. Einsatzkonditionen (Moved to Page 1)
    konditionen = data.get("einsatzkonditionen", {})
    doc.add_heading("Einsatzkonditionen", level=1)
    doc.add_paragraph("Dieses Angebot basiert auf folgenden Konditionen:")
    
    # Load table styles
    table_cfg = styles_cfg.get("table", {})
    width_total = table_cfg.get("width_cm", 17.0)
    ratios = table_cfg.get("column_ratios_conditions", [0.20, 0.80])
    f_size = table_cfg.get("font_size", 10)

    table = doc.add_table(rows=5, cols=2)
    table.style = None
    table.autofit = False
    table.width = Cm(width_total)
    
    # Set column widths & row height
    widths = [Cm(width_total * r) for r in ratios]
    for row in table.rows:
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for i, width in enumerate(widths):
            row.cells[i].width = width
            set_cell_padding(row.cells[i])
            remove_cell_borders(row.cells[i])
    
    # Validity date calculation
    validity_date = get_default_validity_date().strftime("%d.%m.%Y")
    
    rows_cond = [
        ("Pensum:", konditionen.get("pensum", "")),
        ("Verfügbarkeit:", konditionen.get("verfuegbarkeit", "")),
        ("Stundensatz (CHF):", konditionen.get("stundensatz", "")),
        ("Subunternehmen:", konditionen.get("subunternehmen", "")),
        ("Angebot Gültigkeit:", f"{validity_date} ! bitte prüfen !")
    ]
    
    for i, (label, value) in enumerate(rows_cond):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(f_size)
        
        row.cells[1].text = value
        run_v = row.cells[1].paragraphs[0].runs[0]
        run_v.font.size = Pt(f_size)
        # Highlight all condition values or if they contain placeholder
        if "<" in value or "bitte prüfen" in value.lower():
            run_v.font.highlight_color = WD_COLOR_INDEX.YELLOW
        set_cell_background(row.cells[1], "FFFF00") # Yellow

    # 2. Kriterien Abgleich
    doc.add_page_break()
    abgleich = data.get("kriterien_abgleich", {})
    doc.add_heading("Abgleich mit Anforderungskriterien", level=1)
    doc.add_paragraph("Die Muss- und Soll-Kriterien werden erfüllt, mit folgenden ergänzenden Hinweisen:")
    
    # Muss-Kriterien
    add_criteria_table(doc, "Muss-Kriterien", abgleich.get("muss_kriterien", []))

    # Soll-Kriterien
    add_criteria_table(doc, "Soll-Kriterien", abgleich.get("soll_kriterien", []))

    # Weitere Kriterien
    add_criteria_table(doc, "Weitere Anforderungen", abgleich.get("weitere_kriterien", []))

    # Soft Skills
    add_criteria_table(doc, "Persönliche Kompetenzen / Soft Skills", abgleich.get("soft_skills", []))

    # 3. Gesamtbeurteilung
    doc.add_page_break()
    beurteilung = data.get("gesamtbeurteilung", {})
    doc.add_heading("Gesamtbeurteilung & Mehrwert", level=1)
    add_paragraph_with_bold(doc, beurteilung.get("zusammenfassung", ""))
    
    doc.add_heading("Mehrwert für den Kunden", level=2)
    for item in beurteilung.get("mehrwert_fuer_kunden", []):
        add_bullet_paragraph(doc, item)
        
    p_empfehlung = add_paragraph_with_bold(doc, beurteilung.get("empfehlung", ""))
    for run in p_empfehlung.runs:
        run.bold = True

    # Abschluss (No header)
    doc.add_paragraph("\nFür Fragen zu diesem Angebot stehen wir jederzeit gerne zur Verfügung.")

    doc.add_paragraph("\nBeste Grüsse\n")
    add_paragraph_with_bold(doc, "Vorname Nachname ! bitte prüfen !", style=None)
    p_closing = doc.add_paragraph()
    p_closing.add_run("Orange Business Digital").bold = True

    # Save
    doc.save(output_path)
    print(f"✅ Angebot Word generiert: {output_path}")
    return output_path

if __name__ == "__main__":
    # Test
    import sys
    if len(sys.argv) > 2:
        generate_angebot_word(sys.argv[1], sys.argv[2])
