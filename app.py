import streamlit as st
import json
from docx import Document

st.set_page_config(page_title="CV Generator", page_icon="📄")

st.title("📄 CV Generator – JSON → Word")

# --- Hilfsfunktion ---
def list_to_text(lst):
    if not lst:
        return ""
    return "\n".join([f"• {item}" for item in lst])

# --- File Upload ---
uploaded_json = st.file_uploader("JSON-Datei hochladen", type=["json"])

if uploaded_json:
    data = json.load(uploaded_json)
    st.success("JSON erfolgreich geladen!")

    # Template laden
    doc = Document("template.docx")

    # Platzhalter ersetzen
    replacements = {
        "{{Vorname}}": data.get("Vorname", ""),
        "{{Name}}": data.get("Name", ""),
        "{{Hauptrolle}}": data.get("Hauptrolle", ""),
        "{{Nationalität}}": data.get("Nationalität", ""),
        "{{Hauptausbildung}}": data.get("Hauptausbildung", ""),
        "{{Kurzprofil}}": data.get("Kurzprofil", ""),
        "{{Fachwissen_und_Schwerpunkte}}": list_to_text(data.get("Fachwissen_und_Schwerpunkte", [])),
        "{{Aus_und_Weiterbildung}}": list_to_text(data.get("Aus_und_Weiterbildung", [])),
        "{{Trainings_und_Zertifizierungen}}": list_to_text(data.get("Trainings_und_Zertifizierungen", [])),
        "{{Sprachen}}": list_to_text(data.get("Sprachen", [])),
        "{{Referenzprojekte}}": list_to_text(data.get("Referenzprojekte", []))
    }

    # Text ersetzen
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)

    # Auch Tabellen ersetzen (falls Template später Tabellen nutzt)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    # Word als Download bereitstellen
    output_path = "output.docx"
    doc.save(output_path)

    with open(output_path, "rb") as f:
        st.download_button("📥 Word-Dokument herunterladen", f, file_name="CV_Generiert.docx")
